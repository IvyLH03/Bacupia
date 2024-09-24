from typing import List
import requests
import json
from docx import Document
import re
import html
import time
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from zoneinfo import ZoneInfo
import datetime
import copy

class SaveThread:
    def __init__(self, tid: int, authorid: int, cookies):
        self.tid = tid

        self.client = requests.session()
        self.client.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:91.0) Gecko/20100101 Firefox/91.0',
        }
        self.client.cookies.update(cookies)

        self.authorid = authorid
        self.max_page = self.get_thread_pgnum()

    # 获取一页帖子的json
    def get_page(self, pgnum: int):
        payload = {
            "tid":self.tid,
            "page":pgnum,
            "authorid":self.authorid
        }
        res = self.client.post("https://bbs.nga.cn/app_api.php?__lib=post&__act=list", data=payload)
        return res.json()

    # 获取一个帖子的页数
    def get_thread_pgnum(self):
        return self.get_page(1)['totalPage']

    # 获取一页帖子的回复列表
    def get_page_posts(self, pgnum: int) -> List[object]:
        return self.get_page(pgnum)['result']

    # 获取整个帖子的回复列表
    def get_thread_posts(self) -> List[object]:
        pgnum = self.get_thread_pgnum()
        posts = []
        for i in range(1, self.max_page + 1):
            print("fetching posts from page", i, "/", self.max_page, "...")
            posts += (self.get_page_posts(i))
        return posts
            

    # 处理一条回复的文字部分, 生成用于排版的纯文字
    def process_post_content_minimal(self, content: str) -> str:
        # 删除html标签
        content = self.reduce_html(content)

        # 删除特定方括号标签
        content = self.reduce_square_labels(content)

        # 跳过回复楼层
        if content.__contains__("Reply to") or content.__contains__("Reply[/pid] Post by "):
            #print("reply skipped")
            return ""

        # 检测重复d2
        content_list = content.split("\n")
        # ROLL : d10
        # ROLL : d2
        for i in range(0, len(content_list) - 1):
            prev_item = content_list[i]
            cur_item = content_list[i+1]
            # 如果第i条是d10且出目不是10，同时第i+1条是d2，则删除d2
            if prev_item.startswith("ROLL : d10") and cur_item.startswith("ROLL : d2") and (not prev_item.startswith("ROLL : d10=d10(10)=10")):
                content_list[i+1] = ""

        # 重新组合结果
        content = ""
        for i in range(0, len(content_list)):
            if len(content_list[i]) > 0:
                content += content_list[i] + "\n"

        return content.strip()
    
    # 根据html换行符切割帖子内容并保留换行符
    def split_content_keep_newline(self, content:str) -> List[str]:
        return re.split('(<br/>)', content)
    
    def split_content(self, content:str) -> List[str]:
        return content.split('<br/>')
    
    
    # 保存所有回帖的json数据文件
    def save_raw(self, posts:List[object], filename:str):
        print("saving raw to", filename, "...")
        with open(filename,"w",encoding="utf-8") as f:
            json.dump(posts, f)
        print("raw saved!")

    # 保存经过字符串处理的json数据文件
    def save_processed(self, posts:List[object], filename:str):
        print("processing post string contents...")
        post_processed = []

        for post in posts:
            content = str(post['content'])
            content = self.reduce_html(content)
            content = self.reduce_square_labels(content)

             # 检测重复d2
            content_list = content.split("\n")
            # ROLL : d10
            # ROLL : d2
            for i in range(0, len(content_list) - 1):
                prev_item = content_list[i]
                cur_item = content_list[i+1]
                # 如果第i条是d10且出目不是10，同时第i+1条是d2，则删除d2
                if prev_item.startswith("ROLL : d10") and cur_item.startswith("ROLL : d2") and (not prev_item.startswith("ROLL : d10=d10(10)=10")):
                    content_list[i+1] = ""


            post_processed.append({'lou':post['lou'], 'content':content_list})

        with open(filename, "w", encoding="utf-8") as f:
            json.dump(post_processed, f)
        

            
    # 保存所有回帖的排版文档
    def save_minimal(self, posts:List[object], filename:str):
        print("saving minimal to", filename, "...")
        document = Document()
        style = document.styles.add_style('ThreadMinimal', WD_STYLE_TYPE.PARAGRAPH)
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(0)

        for post in posts:
            content = self.process_post_content_minimal(str(post['content']))
            if len(content) > 0:
                paragraph = document.add_paragraph(content)
                paragraph.style = document.styles['ThreadMinimal']
        
        document.save(filename)
        print("minimal saved!")

    # 保存所有回帖的阅读格式文档
    def save_reading(self, posts:List[object], filename:str):
        print("saving reading to", filename, "...")
        document = Document()
        style = document.styles.add_style('ThreadMinimal', WD_STYLE_TYPE.PARAGRAPH)
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(0)


        for post in posts:
            # print('Saving #',post['lou'])
            paragraph = document.add_paragraph()
            run = paragraph.add_run("#" + str(post['lou']) + " [" + post_time_formatted(post['postdatetimestamp']) + "]\n")
            run.add_break()
            paragraph.style = document.styles['ThreadMinimal']

            content_list = self.split_content_keep_newline(post['content'])
            # print(content_list)

            for raw_content in content_list:
                content = self.reduce_html(raw_content)
                content = self.reduce_square_labels(content)
                is_dice_block = self.check_dice_block(raw_content)

                if(content == '\n'):
                    run.add_break()

                elif(is_dice_block):
                    run = paragraph.add_run(content)
                    run.bold = True
                    run = paragraph.add_run()
                else:
                    run.add_text(content)


            # TODO: detect other styles
            # TODO: detect styles within one line
            run.add_break()
            run.add_break()
        document.save(filename)
        print("reading saved!")

            


    # 去除html代码
    def reduce_html(self, content: str) -> str:
        # 变更换行符
        content = content.replace("<br/>","\n")

        # html tags
        content = re.sub('<[^<]+?>', '', content)

        # html entities
        content = html.unescape(content)

        # 替换unicode
        def subchr(s):
            return chr(int(s.group(1)))
        content = re.sub('&#([0-9]+);',subchr,content)
        return content
    
    # 删除特定方括号标签
    def reduce_square_labels(self, content:str) -> str:
        labels = ['[i]','[/i]','[color=red]', '[/color]']
        for label in labels:
            content = content.replace(label, '')
        return content

    # 检测是否是骰点块
    def check_dice_block(self, content: str) -> bool:
        return content.startswith('<div class=\'dice\'>')
        return re.match('<div class=\'dice\'>(.)+</div>', content) != None
    
    # 从已有数据文件生成文档
    def run_save_from_json(self, thread_name:str, json_path:str, save_minimal=True, save_reading=True):
        time_suffix = time.strftime("%Y%m%d_%H%M%S")
        with open(json_path) as f:
            posts = json.load(f)
        if save_minimal:
            filename = thread_name + "_无格式_" + time_suffix + ".docx"
            self.save_minimal(posts, filename)
        if save_reading:
            filename = thread_name + "_阅读版_" + time_suffix + ".docx"
            self.save_reading(posts, filename)

    # 生成文档
    def run_save(self, thread_name:str, save_raw=True, save_minimal=True, save_reading=True):
        posts = self.get_thread_posts()
        time_suffix = time.strftime("%Y%m%d_%H%M%S")
        if save_raw:
            filename = thread_name + "_raw_" + time_suffix + ".json"
            self.save_raw(posts, filename)
        if save_minimal:
            filename = thread_name + "_无格式_" + time_suffix + ".docx"
            self.save_minimal(posts, filename)
        if save_reading:
            filename = thread_name + "_阅读版_" + time_suffix + ".docx"
            self.save_reading(posts, filename)



def post_time_formatted(tstamp:int) -> str:
    target_timezone = ZoneInfo("Asia/Shanghai")
    return datetime.datetime.fromtimestamp(tstamp).astimezone(target_timezone).strftime("%Y-%m-%d %H:%M:%S")

if __name__ == "__main__":
    with open('config.json') as f:
        cookies=json.load(f)
    saver = SaveThread(40452148, 64875447, cookies)
    file_name = "./out/processed_test_0.json"
    #saver.run_save(thread_name=thread_name)
    with open('saves/百命海猎_raw_20240924_124555.json') as f:
        posts = json.load(f)
    saver.save_processed(posts, file_name)